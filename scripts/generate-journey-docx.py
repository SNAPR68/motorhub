#!/usr/bin/env python3
"""Generate CaroBest User Journey Word Document with visual diagrams."""

import io
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------- helpers ----------

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="1a1a2e"):
    """Add a formatted table with colored header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "f0f0f5")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph("")
    return table

def add_diagram_box(doc, title, lines, width=5.5):
    """Add a visual block diagram as a styled table (single-cell with border)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Inches(width)

    # Title
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Content
    for line in lines:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.size = Pt(8)
        run.font.name = "Courier New"
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)

    set_cell_shading(cell, "f8f9fc")

    # Border
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="8" w:color="1a1a2e"/>'
        '<w:left w:val="single" w:sz="8" w:color="1a1a2e"/>'
        '<w:bottom w:val="single" w:sz="8" w:color="1a1a2e"/>'
        '<w:right w:val="single" w:sz="8" w:color="1a1a2e"/>'
        '</w:tcBorders>'
    )
    tc_pr.append(borders)

    doc.add_paragraph("")

def add_flow_diagram(doc, title, steps, arrow="-->"):
    """Add a horizontal flow diagram as a table row."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    cols = len(steps) * 2 - 1  # steps + arrows between
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, step in enumerate(steps):
        cell_idx = i * 2
        cell = table.rows[0].cells[cell_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(step)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1a1a2e")

        # Arrow cell
        if i < len(steps) - 1:
            arrow_cell = table.rows[0].cells[cell_idx + 1]
            arrow_cell.text = ""
            ap = arrow_cell.paragraphs[0]
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ar = ap.add_run("  -->  ")
            ar.font.size = Pt(9)
            ar.bold = True
            ar.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    doc.add_paragraph("")

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_page_break(doc):
    doc.add_page_break()

# ---------- main document ----------

def generate():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # ===========================
    # COVER PAGE
    # ===========================
    for _ in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CAROBEST")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Complete User Journey Document")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x6a)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI-Powered Used Car Marketplace for India")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x6a, 0x6a, 0x8a)

    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("For: Product, Engineering, Marketing & Content Teams")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x8a, 0x8a, 0xaa)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Purpose: Team Education + Explainer Video Scripting")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x8a, 0x8a, 0xaa)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("March 2026  |  Version 1.0")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x8a, 0x8a, 0xaa)

    add_page_break(doc)

    # ===========================
    # TABLE OF CONTENTS
    # ===========================
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Platform Overview",
        "2. System Architecture",
        "3. Buyer Journey (9 Phases)",
        "   3.1 Entry & Discovery",
        "   3.2 Vehicle Deep Dive",
        "   3.2a 360-Degree Virtual Tour",
        "   3.3 Comparison & Decision",
        "   3.4 Engagement & Booking",
        "   3.5 Finance & Insurance Tools",
        "   3.6 Trade-In & Services",
        "   3.7 Post-Purchase Hub",
        "   3.8 AI Concierge",
        "   3.9 SEO Marketplace Pages",
        "4. Dealer Journey (6 Phases)",
        "   4.1 Onboarding & Signup",
        "   4.2 Dashboard & Daily Ops",
        "   4.3 Lead Management",
        "   4.4 AI Studio",
        "   4.5 Analytics & Intelligence",
        "   4.6 Settings & Preferences",
        "5. Data Flow Architecture",
        "6. AI Pipeline",
        "7. Event System & Agent Infrastructure",
        "8. Payment Flow",
        "9. Error Handling Architecture",
        "Appendix: Notification System & Auth Flow",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)

    add_page_break(doc)

    # ===========================
    # 1. PLATFORM OVERVIEW
    # ===========================
    add_heading(doc, "1. Platform Overview", level=1)

    add_body(doc, "CaroBest is an AI-powered used car marketplace for India serving two distinct user types:")

    add_bullet(doc, " Browse, compare, finance, and purchase used cars", bold_prefix="Buyers (Consumers):")
    add_bullet(doc, " List inventory, manage leads, use AI tools, track analytics", bold_prefix="Dealers (Businesses):")

    add_body(doc, "Business Model: Dealers pay monthly subscriptions (STARTER / GROWTH / ENTERPRISE) for access to DealerOS tools. Buyers use the platform for free.")

    add_diagram_box(doc, "PLATFORM ARCHITECTURE", [
        "+-----------------------+     +-----------------------+",
        "|     BUYER SIDE        |     |     DEALER SIDE       |",
        "|                       |     |                       |",
        "|  Browse & Search      |     |  List Vehicles        |",
        "|  Compare Vehicles     |     |  Manage Leads         |",
        "|  Get Finance Info     |     |  AI Photo Studio      |",
        "|  Book Services        |     |  AI Content Tools     |",
        "|  AI Concierge Chat    |     |  Analytics Dashboard  |",
        "|  Track Purchases      |     |  Market Intelligence  |",
        "+-----------+-----------+     +-----------+-----------+",
        "            |                             |            ",
        "            +-------------+---------------+            ",
        "                          |                             ",
        "              +-----------v-----------+                 ",
        "              |   SHARED PLATFORM     |                 ",
        "              |                       |                 ",
        "              |  Supabase Auth + DB   |                 ",
        "              |  Prisma ORM           |                 ",
        "              |  OpenAI + Replicate   |                 ",
        "              |  Event System         |                 ",
        "              |  Razorpay Payments    |                 ",
        "              |  Sentry Monitoring    |                 ",
        "              +-----------------------+                 ",
    ])

    add_heading(doc, "Platform Stats", level=2)
    add_styled_table(doc,
        ["Metric", "Count"],
        [
            ["Total Routes (Pages)", "204+"],
            ["API Endpoints", "40+"],
            ["AI Features", "12"],
            ["Database Models", "20+"],
            ["Stitch Designs Implemented", "75+"],
            ["Killer Features", "10"],
            ["Static SEO Pages", "167"],
            ["Car Brands in DB", "12"],
            ["Car Models in DB", "17"],
            ["Car Variants in DB", "24"],
        ],
        col_widths=[3.5, 2.0]
    )

    add_page_break(doc)

    # ===========================
    # 2. SYSTEM ARCHITECTURE
    # ===========================
    add_heading(doc, "2. System Architecture", level=1)

    add_heading(doc, "Tech Stack", level=2)
    add_styled_table(doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "Next.js 16 + React 19", "Server-rendered pages + client interactivity"],
            ["Styling", "Tailwind CSS v4", "Utility-first responsive design"],
            ["State", "Zustand 5", "Client-side state management"],
            ["Animation", "Framer Motion 12", "Page transitions + micro-interactions"],
            ["Database", "Supabase PostgreSQL", "Managed Postgres with realtime"],
            ["ORM", "Prisma 7.4", "Type-safe database queries"],
            ["Auth", "Supabase Auth", "Email/password + Google OAuth + Magic Link"],
            ["AI (Text)", "OpenAI GPT-4o / GPT-4o-mini", "Descriptions, sentiment, replies, scripts"],
            ["AI (Image)", "Replicate (Stable Diffusion, rembg)", "Photo editing, background removal"],
            ["Payments", "Razorpay", "Subscription billing (INR)"],
            ["Monitoring", "Sentry", "Error tracking + performance"],
            ["Analytics", "PostHog", "Product analytics + events"],
            ["Hosting", "Vercel", "Serverless deployment + CDN"],
        ],
        col_widths=[1.2, 2.3, 2.5]
    )

    add_heading(doc, "Request Flow", level=2)
    add_flow_diagram(doc, "How Data Flows Through the App",
        ["Browser", "useApi()", "fetch()", "API Route", "Prisma", "Database"])

    add_diagram_box(doc, "DETAILED REQUEST LIFECYCLE", [
        "1. User interacts with UI component (React 19)",
        "2. Component calls useApi(fetchFunction, dependencies)",
        "3. useApi manages loading/error/data states automatically",
        "4. fetchFunction calls fetch('/api/...') with auth headers",
        "5. Next.js API Route receives request",
        "6. Route checks auth via Supabase getUser()",
        "7. Route validates input with Zod (POST/PATCH)",
        "8. Route executes business logic via Prisma ORM",
        "9. Prisma queries Supabase PostgreSQL",
        "10. Response returned as JSON",
        "11. useApi updates component state",
        "12. UI re-renders with new data",
    ])

    add_page_break(doc)

    # ===========================
    # 3. BUYER JOURNEY
    # ===========================
    add_heading(doc, "3. Buyer Journey", level=1)

    add_body(doc, "The buyer journey spans 9 phases, from initial discovery through post-purchase ownership. Each phase is designed to reduce friction and build trust through AI-powered features.")

    add_flow_diagram(doc, "Buyer Journey Overview",
        ["Discover", "Browse", "Deep Dive", "Compare", "Engage", "Finance", "Services", "Purchase", "Own"])

    # 3.1
    add_heading(doc, "3.1 Entry & Discovery", level=2)

    add_heading(doc, "Landing Page (/)", level=3)
    add_bullet(doc, "Hero section with search bar (brand / model / budget)")
    add_bullet(doc, "Featured vehicles carousel showing top-rated cars")
    add_bullet(doc, "Trust signals: total vehicles, verified dealers, cities covered")
    add_bullet(doc, "Quick-access links to popular brands (Hyundai, Maruti, Tata, Honda)")
    add_bullet(doc, "Quick-access links to body types (SUV, Sedan, Hatchback)")

    add_heading(doc, "Showroom (/showroom)", level=3)
    add_bullet(doc, "Grid of vehicle cards showing: image, name, price, year, km, location, AI score")
    add_bullet(doc, "Filter panel: brand, model, budget range, fuel type, body type, year range")
    add_bullet(doc, "Sort options: price (low-high, high-low), newest first, most popular")
    add_bullet(doc, "Infinite scroll pagination for seamless browsing")
    add_bullet(doc, "Each card displays AI Quality Score badge (0-100)")

    add_diagram_box(doc, "SHOWROOM CARD LAYOUT", [
        "+---------------------------+",
        "|       [Vehicle Image]     |",
        "|                           |",
        "|  Hyundai Creta SX 2022   |",
        "|  8.5 Lakh  |  45,000 km  |",
        "|  Diesel  |  Automatic     |",
        "|                           |",
        "|  [AI Score: 85]  Delhi    |",
        "|                           |",
        "|  [Wishlist]  [Contact]    |",
        "+---------------------------+",
    ])

    # 3.2
    add_heading(doc, "3.2 Vehicle Deep Dive", level=2)

    add_heading(doc, "Vehicle Detail Page (/vehicle/[id])", level=3)
    add_bullet(doc, "Full-screen swipeable image gallery (up to 20 photos)")
    add_bullet(doc, "Price display with EMI estimate")
    add_bullet(doc, "AI-generated vehicle description (via OpenAI)")
    add_bullet(doc, "Specifications: fuel, transmission, engine, odometer, owners, insurance")
    add_bullet(doc, "Dealer info card with average response time")
    add_bullet(doc, "CTA buttons: 'Contact Dealer' (creates Lead) and 'Add to Wishlist'")
    add_bullet(doc, "Tabs: Overview, Specifications, Similar Cars")

    add_heading(doc, "Vehicle Passport (/vehicle/passport/[id])", level=3)
    add_body(doc, "CaroBest's proprietary trust report. A unique selling point that builds buyer confidence.")

    add_styled_table(doc,
        ["Check", "What It Shows", "Rating"],
        [
            ["Ownership History", "Number of previous owners, transfer dates", "Green / Yellow / Red"],
            ["Service Records", "Service regularity, authorized center visits", "Green / Yellow / Red"],
            ["Accident Check", "Structural damage, insurance claims", "Green / Yellow / Red"],
            ["Flood Check", "Water damage indicators", "Green / Yellow / Red"],
            ["Theft Check", "Police records, stolen vehicle database", "Green / Yellow / Red"],
            ["Insurance Status", "Current policy, claim history", "Green / Yellow / Red"],
        ],
        col_widths=[1.5, 3.0, 1.5]
    )

    add_bullet(doc, "QR code on each passport for easy sharing")
    add_bullet(doc, "Deterministically generated from vehicle data (no external API dependency)")

    add_heading(doc, "TrueCost Calculator (/vehicle/[id]/true-cost)", level=3)
    add_body(doc, "Shows the TOTAL cost of ownership, not just the sticker price. Computed client-side.")

    add_styled_table(doc,
        ["Cost Category", "What's Included"],
        [
            ["Registration & Transfer", "RTO fees, state transfer tax, stamp duty"],
            ["Insurance", "1st year premium + 3-year renewal estimate"],
            ["Maintenance", "Annual servicing based on brand/model averages"],
            ["Fuel Cost", "Annual estimate based on average driving patterns"],
            ["Depreciation", "3-year projected value loss"],
        ],
        col_widths=[2.0, 4.0]
    )

    add_heading(doc, "AI Negotiation Coach (/vehicle/[id]/negotiate)", level=3)
    add_bullet(doc, "Client-side AI that helps buyers negotiate better")
    add_bullet(doc, "Input: listed price, target price, vehicle condition notes")
    add_bullet(doc, "Output: negotiation strategy, talking points, fair price range")
    add_bullet(doc, "Uses market data + vehicle AI score for recommendations")

    add_page_break(doc)

    # 3.2a — 360-Degree Virtual Tour
    add_heading(doc, "3.2a: 360-Degree Virtual Tour", level=2)

    add_body(doc, "The Virtual Tour is one of CaroBest's most immersive buyer experiences. It allows buyers to explore vehicles from every angle through drag-to-rotate interaction, gyroscope control, and interactive hotspots — all powered by real vehicle data from the API.")

    add_flow_diagram(doc, "Virtual Tour End-to-End Flow",
        ["Dealer\nUploads", "Marks\nPanorama", "Buyer Opens\nTour", "Drag to\nRotate", "View\nHotspots", "Book\nVisit"])

    add_heading(doc, "Main Virtual Tour (/virtual-tour/[id])", level=3)
    add_body(doc, "The primary 360-degree experience. A drag-to-rotate panorama viewer with physics-based momentum scrolling.")

    add_diagram_box(doc, "VIRTUAL TOUR SCREEN LAYOUT", [
        "+---------------------------------------------------+",
        "|  [Back]  Hyundai Creta SX        360 VIEW  [Gyro] |",
        "|          ============================              |",
        "|                                                    |",
        "|  +----------------------------------------------+  |",
        "|  |                                              |  |",
        "|  |        << DRAG TO EXPLORE >>                 |  | o",
        "|  |                                              |  | o  Compass",
        "|  |     [Hotspot: LED DRLs]                      |  | o  Ring",
        "|  |                                              |  | .",
        "|  |              [Hotspot: Alloy Wheels]         |  | .",
        "|  |                                              |  |",
        "|  +----------------------------------------------+  |",
        "|                                                    |",
        "|  [thumb1] [thumb2] [thumb3] [thumb4] [thumb5]      |",
        "|                                                    |",
        "|  [  Full Specs  ]           [  Book Visit  ]       |",
        "+---------------------------------------------------+",
    ])

    add_heading(doc, "Interaction Features", level=3)
    add_styled_table(doc,
        ["Feature", "How It Works"],
        [
            ["Drag-to-Rotate", "Multi-image horizontal strip at 300% viewport width. Pointer events track drag and translate strip position."],
            ["Momentum Scrolling", "Velocity tracking with 0.92x decay per animation frame. Flick gesture carries motion forward naturally."],
            ["Auto-Rotate", "On first load, images slowly pan across automatically. Stops on first user interaction."],
            ["Gyroscope Control", "Toggle button enables DeviceOrientationEvent. Tilting phone pans through images. Works on mobile only."],
            ["Contextual Hotspots", "Each image index has mapped hotspots (label, category, description) that appear as pulsing dots."],
            ["Compass Ring", "Vertical pill buttons on right edge (one per image). Click to jump directly to any angle."],
            ["Thumbnail Strip", "Bottom carousel shows all images. Click to jump. Current image highlighted."],
            ["CRT Scan Lines", "Retro visual effect on page load. Animated horizontal lines fade out after 3 seconds."],
            ["Panorama Start", "Dealer-marked panoramaImageIdx determines which image loads first (auto-scroll with 400ms delay)."],
        ],
        col_widths=[1.5, 4.5]
    )

    add_heading(doc, "Interior Viewer (/virtual-tour/viewer/[id])", level=3)
    add_body(doc, "A secondary viewer focused on interior exploration with a simpler interaction model.")

    add_diagram_box(doc, "INTERIOR VIEWER LAYOUT", [
        "+---------------------------------------------------+",
        "|  [Back]   Interior View              [1/4]        |",
        "|                                                    |",
        "|  [+]  +-------------------------------+  [img1]   |",
        "|  [-]  |                               |  [img2]   |",
        "|       |    [*] Dashboard              |  [img3]   |",
        "|       |                               |  [img4]   |",
        "|       |         [*] Steering          |           |",
        "|       |                               |           |",
        "|       |    [*] Upholstery             |           |",
        "|       |                               |           |",
        "|       +-------------------------------+           |",
        "|                                                    |",
        "|  [Dashboard] [Leather Quality] [Audio] [Ambient]  |",
        "|                                                    |",
        "|  AI Concierge: Have questions? Ask our AI -->      |",
        "+---------------------------------------------------+",
    ])

    add_bullet(doc, "Full-screen background image with gradient vignette overlay")
    add_bullet(doc, "3 interactive hotspots positioned on interior features")
    add_bullet(doc, "Zoom controls (+/-) on left side for detail inspection")
    add_bullet(doc, "Tag pills at bottom for quick feature category browsing")
    add_bullet(doc, "Image sidebar thumbnails (up to 4 visible) for switching views")
    add_bullet(doc, "Links to AI Concierge for follow-up questions")

    add_heading(doc, "Virtual Tour Routes", level=3)
    add_styled_table(doc,
        ["Route", "Purpose", "Data Source", "Interaction"],
        [
            ["/virtual-tour/[id]", "Individual vehicle panorama", "Real API: fetchVehicle()", "Drag-to-rotate + gyroscope"],
            ["/virtual-tour/viewer/[id]", "Interior focused viewer", "Real API: fetchVehicle()", "Click hotspots + zoom"],
            ["/[brand]/[model]/360-view", "Brand/model template", "Static hotspot data", "Toggle exterior/interior + hotspots"],
        ],
        col_widths=[1.8, 1.5, 1.5, 1.5]
    )

    add_heading(doc, "Dealer: Panorama Management", level=3)
    add_body(doc, "Dealers control which image appears first in the virtual tour by marking a 'panorama hero' angle.")

    add_diagram_box(doc, "HOW PANORAMA WORKS (END-TO-END)", [
        "1. Dealer uploads vehicle with multiple photos (up to 20)",
        "2. Photos stored in Vehicle.images[] array (Supabase Storage)",
        "3. Dealer opens inventory, selects vehicle",
        "4. Marks best angle as panorama image",
        "5. PATCH /api/vehicles/[id]/panorama",
        "   --> { panoramaImageIdx: 2 }",
        "   --> Saved to Vehicle.panoramaImageIdx in DB",
        "",
        "6. Buyer opens /virtual-tour/[vehicleId]",
        "7. Frontend fetches vehicle from API",
        "   --> Gets images: [url0, url1, url2, ...] + panoramaImageIdx: 2",
        "8. Tour auto-scrolls to image index 2 on load (400ms delay)",
        "9. User drags left/right to pan through all images",
        "10. Hotspots appear contextually per image",
        "11. CTAs: 'Full Specs' or 'Book Visit' to continue journey",
    ])

    add_styled_table(doc,
        ["Component", "Detail"],
        [
            ["API Endpoint", "PATCH /api/vehicles/[id]/panorama"],
            ["Request Body", "{ panoramaImageIdx: number | null }"],
            ["Client Function", "setPanoramaImage(vehicleId, index) in src/lib/api.ts"],
            ["DB Field", "Vehicle.panoramaImageIdx (optional Int)"],
            ["Validation", "Must be non-negative integer or null; requires dealer auth + ownership check"],
        ],
        col_widths=[1.5, 4.5]
    )

    add_page_break(doc)

    # 3.3
    add_heading(doc, "3.3 Comparison & Decision", level=2)

    add_heading(doc, "Compare Page (/compare)", level=3)
    add_body(doc, "Side-by-side comparison of 2-3 vehicles on all key parameters:")

    add_styled_table(doc,
        ["Parameter", "Vehicle A", "Vehicle B", "Vehicle C"],
        [
            ["Price", "8.5 Lakh", "7.2 Lakh", "9.1 Lakh"],
            ["Year", "2022", "2023", "2022"],
            ["Odometer", "45,000 km", "28,000 km", "38,000 km"],
            ["AI Score", "85", "92 (Winner)", "78"],
            ["Engine", "1.5L Diesel", "1.2L Petrol", "1.5L Petrol"],
            ["Power", "115 HP", "120 HP", "121 HP"],
            ["Mileage", "21 kmpl", "20 kmpl", "18 kmpl"],
            ["Safety Rating", "5 Star", "5 Star", "4 Star"],
        ],
        col_widths=[1.5, 1.5, 1.5, 1.5]
    )

    add_heading(doc, "Wishlist (/wishlist)", level=3)
    add_bullet(doc, "Saved vehicles with quick-access actions")
    add_bullet(doc, "Price drop alerts when wishlisted cars get cheaper")
    add_bullet(doc, "Quick actions: compare, contact dealer, remove")
    add_bullet(doc, "When a vehicle gets 3+ wishlists, it gets a 'Trending' badge (AI agent-driven)")

    # 3.4
    add_heading(doc, "3.4 Engagement & Booking", level=2)

    add_flow_diagram(doc, "Contact-to-Conversation Flow",
        ["Buyer Clicks\n'Contact'", "Lead\nCreated", "AI Analyzes\nSentiment", "Auto-Reply\nSent", "Dealer\nNotified", "Conversation\nBegins"])

    add_diagram_box(doc, "WHAT HAPPENS WHEN A BUYER CLICKS 'CONTACT DEALER'", [
        "1. Lead record created in database (POST /api/leads)",
        "2. PlatformEvent emitted: LEAD_CREATED",
        "3. AI Agent auto-analyzes buyer sentiment (HOT/WARM/COOL)",
        "4. If dealer opted in: AI generates and sends auto-reply",
        "5. Follow-up reminders scheduled (24h, 48h, 72h)",
        "6. Dealer receives push notification",
        "7. Conversation thread opens for buyer-dealer messaging",
        "",
        "All steps 2-6 happen in background (fire-and-forget).",
        "Buyer sees confirmation instantly.",
    ])

    # 3.5
    add_heading(doc, "3.5 Finance & Insurance Tools", level=2)

    add_styled_table(doc,
        ["Tool", "Route", "What It Does"],
        [
            ["EMI Calculator", "/car-loan/emi-calculator", "Monthly payment calculation with amortization schedule"],
            ["Loan Eligibility", "/car-loan/eligibility", "Check loan eligibility based on income and credit"],
            ["Insurance Quotes", "/car-insurance/*", "Compare comprehensive, third-party, and standalone OD"],
            ["Fuel Price", "/fuel-price", "City-wise petrol, diesel, CNG prices with trends"],
            ["RTO Tools", "/rto", "Registration fee calculator + documents checklist"],
        ],
        col_widths=[1.5, 2.0, 2.5]
    )

    # 3.6
    add_heading(doc, "3.6 Trade-In & Services", level=2)

    add_heading(doc, "SwapDirect (/swap)", level=3)
    add_flow_diagram(doc, "Trade-In Flow",
        ["Enter Car\nDetails", "AI\nValuation", "See Trade\nValue", "Browse\nMatches", "Book\nSwap Deal"])

    add_bullet(doc, "Enter current car details (brand, model, year, km, condition)")
    add_bullet(doc, "AI-powered instant valuation via /api/ai/valuation")
    add_bullet(doc, "Browse matching vehicles within budget after swap credit")
    add_bullet(doc, "Book swap deal (creates ServiceBooking)")

    add_heading(doc, "Other Services", level=3)
    add_styled_table(doc,
        ["Service", "Route", "Description"],
        [
            ["InstantRC Transfer", "/rc-transfer", "RC ownership transfer service with document tracking"],
            ["LiveCondition Inspection", "/inspection", "200-point professional inspection booking"],
            ["CrossState Express", "/cross-state", "Cross-state vehicle transfer with NOC + re-registration"],
        ],
        col_widths=[1.5, 1.5, 3.0]
    )

    add_page_break(doc)

    # 3.7
    add_heading(doc, "3.7 Post-Purchase Hub", level=2)

    add_body(doc, "After purchase, buyers access their ownership hub at /my-account/*:")

    add_styled_table(doc,
        ["Section", "Route", "Features"],
        [
            ["My Garage", "/my-account/garage", "Owned vehicles, service reminders, insurance renewal alerts"],
            ["Warranty", "/my-account/warranty", "Active warranty cards, claim history, coverage details"],
            ["Documents", "/my-account/documents", "RC copy, insurance, PUC uploads and downloads"],
            ["Orders", "/my-account/orders", "Purchase history, service booking tracking"],
            ["Resale", "/my-account/resale", "List car for resale, AI valuation, market demand insights"],
        ],
        col_widths=[1.2, 2.0, 2.8]
    )

    # 3.8
    add_heading(doc, "3.8 AI Concierge", level=2)
    add_body(doc, "Chat-based AI assistant at /concierge that helps buyers throughout their journey:")
    add_bullet(doc, "Remembers conversation history across sessions (ConciergeMemory table)")
    add_bullet(doc, "Can recommend vehicles based on stated preferences")
    add_bullet(doc, "Answers questions about features, pricing, documentation")
    add_bullet(doc, "Guides buyers through the purchase process step by step")
    add_bullet(doc, "Powered by OpenAI GPT-4o with full conversation context")

    # 3.9
    add_heading(doc, "3.9 SEO Marketplace Pages", level=2)
    add_body(doc, "167 statically-generated pages optimized for search engines:")

    add_styled_table(doc,
        ["Route Pattern", "Example", "Content"],
        [
            ["/used-cars/[city]", "/used-cars/delhi", "Used cars listings for a city"],
            ["/used-cars/[city]/[brand]", "/used-cars/delhi/hyundai", "Brand-filtered city listings"],
            ["/[brand]", "/hyundai", "All models from a brand"],
            ["/[brand]/[model]", "/hyundai/creta", "Model overview page"],
            ["/[brand]/[model]/price", "/hyundai/creta/price", "Pricing information"],
            ["/[brand]/[model]/specifications", "/hyundai/creta/specifications", "Full spec sheet"],
            ["/[brand]/[model]/[variant]", "/hyundai/creta/sx-o", "Variant detail page"],
            ["/dealers/[city]", "/dealers/mumbai", "Dealer directory for a city"],
            ["/car-news", "/car-news", "News articles listing"],
            ["/car-news/[slug]", "/car-news/best-suv-2026", "Individual article"],
        ],
        col_widths=[2.0, 2.0, 2.0]
    )

    add_heading(doc, "Buyer Navigation", level=2)
    add_diagram_box(doc, "BOTTOM NAVIGATION BAR (Mobile)", [
        "+----------+----------+----------+----------+----------+",
        "|   Home   | Showroom | Wishlist |Concierge | Account  |",
        "|  (home)  | (store)  | (heart)  |  (chat)  | (person) |",
        "+----------+----------+----------+----------+----------+",
    ])

    add_page_break(doc)

    # ===========================
    # 4. DEALER JOURNEY
    # ===========================
    add_heading(doc, "4. Dealer Journey", level=1)

    add_body(doc, "The dealer journey spans 6 phases, from onboarding through advanced intelligence tools. DealerOS is designed as a complete business operating system.")

    add_flow_diagram(doc, "Dealer Journey Overview",
        ["Signup", "Dashboard", "Inventory", "Leads", "AI Studio", "Intelligence"])

    # 4.1
    add_heading(doc, "4.1 Onboarding & Signup", level=2)

    add_heading(doc, "5-Step Dealer Registration (/dealer/signup)", level=3)
    add_flow_diagram(doc, "Signup Flow",
        ["Business\nDetails", "Contact\nInfo", "Documents\nUpload", "Showroom\nPhotos", "Plan\nSelection"])

    add_styled_table(doc,
        ["Step", "Fields", "Purpose"],
        [
            ["1. Business Details", "Business name, type, city, state", "Identify the dealership"],
            ["2. Contact Info", "Phone, email, WhatsApp number", "Communication channels"],
            ["3. Documents", "GSTIN, trade license, address proof", "Verification & trust"],
            ["4. Showroom Photos", "Interior, exterior, team photos", "Public profile visuals"],
            ["5. Plan Selection", "STARTER / GROWTH / ENTERPRISE", "Subscription & billing"],
        ],
        col_widths=[1.5, 2.0, 2.5]
    )

    add_heading(doc, "Subscription Plans (/plans)", level=3)
    add_styled_table(doc,
        ["Feature", "STARTER (Free)", "GROWTH (2,999/mo)", "ENTERPRISE (9,999/mo)"],
        [
            ["Vehicle Listings", "5", "50", "Unlimited"],
            ["Lead Management", "Basic", "Full CRM", "Full CRM + API"],
            ["AI Tools", "No", "Yes", "Yes + Custom"],
            ["Analytics", "Basic", "Advanced", "Advanced + Export"],
            ["Support", "Email", "Priority", "Dedicated Manager"],
            ["Team Members", "1", "3", "Unlimited"],
            ["White Label", "No", "No", "Yes"],
        ],
        col_widths=[1.5, 1.5, 1.5, 1.5]
    )

    add_page_break(doc)

    # 4.2
    add_heading(doc, "4.2 Dashboard & Daily Operations", level=2)

    add_heading(doc, "Dashboard (/dashboard)", level=3)
    add_body(doc, "The dealer's command center. Shows everything needed at a glance:")

    add_diagram_box(doc, "DASHBOARD LAYOUT", [
        "+--------------------------------------------------+",
        "|  Welcome, [Dealer Name]          [Notifications] |",
        "|                                                  |",
        "|  +----------+  +----------+  +----------+       |",
        "|  |    24    |  |     8    |  |   2.1h   |       |",
        "|  |  Active  |  |   New    |  |   Avg    |       |",
        "|  |  Cars    |  |  Leads   |  |  Reply   |       |",
        "|  +----------+  +----------+  +----------+       |",
        "|                                                  |",
        "|  +----------+  +----------+                      |",
        "|  |   18%    |  |  12.5L   |                      |",
        "|  | Convert  |  | Revenue  |                      |",
        "|  |  Rate    |  |  Month   |                      |",
        "|  +----------+  +----------+                      |",
        "|                                                  |",
        "|  RECENT ACTIVITY                                 |",
        "|  - New lead: Rahul (Creta) - 2h ago             |",
        "|  - AI auto-replied to Priya - 3h ago            |",
        "|  - Vehicle sold: Swift - yesterday               |",
        "|                                                  |",
        "|  AI AGENT ACTIVITY (Today)                       |",
        "|  - Auto-replies sent: 5                          |",
        "|  - Sentiments analyzed: 12                       |",
        "|  - Follow-ups scheduled: 3                       |",
        "|  - Vehicles auto-scored: 8                       |",
        "+--------------------------------------------------+",
    ])

    add_heading(doc, "Inventory Management (/dashboard/inventory)", level=3)
    add_bullet(doc, "Vehicle list with status badges: AVAILABLE (green), SOLD (purple), RESERVED (yellow)")
    add_bullet(doc, "Bulk actions: mark sold, update price, archive")
    add_bullet(doc, "Each row: image thumbnail, name, price, AI score, views count, leads count")
    add_bullet(doc, "'Add Vehicle' opens full creation form with 5 sections")

    add_heading(doc, "Vehicle Creation Flow", level=3)
    add_flow_diagram(doc, "Adding a New Vehicle",
        ["Basic Info", "Pricing", "Specs", "Photos\nUpload", "AI Auto-\nEnhance"])

    add_diagram_box(doc, "WHAT HAPPENS AFTER 'SAVE VEHICLE'", [
        "1. Vehicle record created in database",
        "2. Photos uploaded to Supabase Storage",
        "3. AI auto-generates description (OpenAI GPT-4o-mini)",
        "4. AI calculates quality score (0-100)",
        "5. PlatformEvent emitted: VEHICLE_CREATED",
        "6. Vehicle appears in showroom for buyers",
        "7. (Optional) Dealer marks panorama hero image",
        "   --> PATCH /api/vehicles/[id]/panorama",
        "   --> Image loads first in buyer's Virtual Tour",
    ])

    add_page_break(doc)

    # 4.3
    add_heading(doc, "4.3 Lead Management", level=2)

    add_heading(doc, "Leads List (/dashboard/leads)", level=3)
    add_body(doc, "All incoming buyer inquiries with comprehensive filtering:")

    add_styled_table(doc,
        ["Filter", "Options"],
        [
            ["Status", "NEW, CONTACTED, QUALIFIED, CLOSED_WON, CLOSED_LOST"],
            ["Sentiment", "HOT (urgent buyer), WARM (interested), COOL (browsing)"],
            ["Source", "WEBSITE, FACEBOOK, INSTAGRAM, WHATSAPP, WALKIN, REFERRAL"],
        ],
        col_widths=[1.5, 4.5]
    )

    add_heading(doc, "Lead Lifecycle", level=3)
    add_diagram_box(doc, "LEAD STATUS PROGRESSION", [
        "               Buyer contacts dealer",
        "                       |",
        "                       v",
        "              +--------+--------+",
        "              |   LEAD CREATED  |   AI: Auto-sentiment",
        "              |   Status: NEW   |   AI: Auto-reply",
        "              +--------+--------+   Schedule follow-ups",
        "                       |",
        "                Dealer responds",
        "                       |",
        "                       v",
        "              +--------+--------+",
        "              |   CONTACTED     |   Record in agent memory",
        "              +--------+--------+   Adjust follow-up",
        "                       |",
        "                Buyer qualifies",
        "                       |",
        "                       v",
        "              +--------+--------+",
        "              |   QUALIFIED     |   Track conversion",
        "              +--------+--------+",
        "                       |",
        "                 +-----+-----+",
        "                 |           |",
        "                 v           v",
        "           +---------+ +---------+",
        "           |CLOSED   | |CLOSED   |",
        "           |WON      | |LOST     |",
        "           |(Deal!)  | |(Lost)   |",
        "           +---------+ +---------+",
    ])

    add_heading(doc, "Lead Detail (/dashboard/leads/[id])", level=3)
    add_bullet(doc, "Full conversation thread showing manual + AI auto messages")
    add_bullet(doc, "Buyer info panel: name, phone, email, interested vehicle")
    add_bullet(doc, "AI Sentiment badge with confidence score (e.g., HOT 87%)")
    add_bullet(doc, "Buyer Intent Signals: keywords detected, urgency indicators")
    add_bullet(doc, "Status change dropdown for quick updates")
    add_bullet(doc, "Message composer for dealer replies")

    add_page_break(doc)

    # 4.4
    add_heading(doc, "4.4 AI Studio", level=2)

    add_body(doc, "A suite of AI-powered tools for creating professional content. Available at /dashboard/studio/*.")

    add_heading(doc, "Photo Studio (/dashboard/studio/photos)", level=3)

    add_styled_table(doc,
        ["Tool", "Technology", "What It Does", "Fallback"],
        [
            ["Background Remove", "Replicate (rembg)", "Removes messy backgrounds from car photos", "Returns original image"],
            ["Mood Application", "Replicate (Stable Diffusion)", "Applies professional lighting styles to photos", "Returns original image"],
        ],
        col_widths=[1.3, 1.5, 1.8, 1.4]
    )

    add_body(doc, "Available Mood Styles:")
    add_styled_table(doc,
        ["Mood", "Visual Effect"],
        [
            ["Golden Hour", "Warm amber tones, long shadows, cinematic feel"],
            ["Midnight", "Deep blue city reflections, dramatic night scene"],
            ["Bloom", "Soft ethereal light, pink and white, dreamy"],
            ["Vivid", "Saturated colors, punchy contrast, high-impact"],
            ["Cinematic", "Teal and orange color grading, film look"],
            ["Matte", "Flat colors, minimal contrast, editorial style"],
        ],
        col_widths=[1.5, 4.5]
    )

    add_heading(doc, "Content Tools", level=3)
    add_styled_table(doc,
        ["Tool", "Route", "Input", "Output"],
        [
            ["Description Generator", "/api/ai/description", "Vehicle details", "SEO-optimized listing description"],
            ["Quick Draft", "/api/ai/quick-draft", "Intent + context", "Marketing copy for any purpose"],
            ["Notification Enhancer", "/api/ai/notification-enhance", "Plain notification text", "Engaging, professional notification"],
            ["Creative Suggestions", "/api/ai/creative/suggestions", "Vehicle + campaign type", "3 marketing campaign ideas"],
            ["Reel Script", "/api/ai/reel/script", "Vehicle + style", "30-second video script"],
            ["Text-to-Speech", "/api/ai/reel/tts", "Script text + voice", "MP3 voiceover audio"],
            ["Smart Reply", "/api/ai/smart-reply", "Lead conversation", "3 contextual reply suggestions"],
        ],
        col_widths=[1.3, 1.5, 1.2, 2.0]
    )

    add_page_break(doc)

    # 4.5
    add_heading(doc, "4.5 Analytics & Intelligence", level=2)

    add_heading(doc, "Performance Analytics (/dashboard/analytics)", level=3)
    add_diagram_box(doc, "ANALYTICS DASHBOARD", [
        "+--------------------------------------------------+",
        "|  PERFORMANCE THIS MONTH                          |",
        "|                                                  |",
        "|  Revenue       New Leads       Conversion Rate   |",
        "|  12.5 Lakh     48              18%               |",
        "|  +23% MoM      +15% MoM        +2% MoM          |",
        "|                                                  |",
        "|  LEAD FUNNEL                                     |",
        "|  New: 48 --> Contacted: 32 --> Qualified: 12     |",
        "|                          --> Closed Won: 8       |",
        "|                          --> Closed Lost: 4      |",
        "|                                                  |",
        "|  TOP VEHICLES (by leads generated)               |",
        "|  1. Hyundai Creta SX 2022 ........ 12 leads     |",
        "|  2. Tata Nexon XZ 2023 ........... 8 leads      |",
        "|  3. Honda City ZX 2022 ........... 6 leads      |",
        "|                                                  |",
        "|  LEAD SOURCE BREAKDOWN                           |",
        "|  Website: 40%   |   WhatsApp: 25%               |",
        "|  Facebook: 20%  |   Walk-in: 15%                |",
        "+--------------------------------------------------+",
    ])

    add_heading(doc, "Dealer Health Score", level=3)
    add_body(doc, "A proprietary composite metric that incentivizes platform engagement. Shown on the dealer's public profile page.")

    add_styled_table(doc,
        ["Factor", "Weight", "What It Measures"],
        [
            ["Response Time", "25%", "Average time to first reply to new leads"],
            ["Conversion Rate", "25%", "Leads closed won / total leads"],
            ["AI Score Average", "20%", "Average quality score of listed vehicles"],
            ["Listing Quality", "15%", "Photo count, description completeness"],
            ["Customer Reviews", "15%", "Average review rating from buyers"],
        ],
        col_widths=[1.5, 0.8, 3.7]
    )

    add_heading(doc, "DemandPulse Intelligence (/intelligence/*)", level=3)
    add_styled_table(doc,
        ["Page", "Route", "Insights Provided"],
        [
            ["Market Overview", "/intelligence", "Overall demand heatmap, trending models"],
            ["Pricing", "/intelligence/pricing", "Real-time competitor pricing, optimal price suggestions"],
            ["Demand Forecast", "/intelligence/forecast", "30-day demand prediction by model"],
            ["Depreciation", "/intelligence/depreciation", "Value loss curves by brand/model/age"],
            ["Best Time to List", "/intelligence/timing", "Optimal listing windows by vehicle type"],
            ["Competitor Intel", "/intelligence/competitors", "Anonymous competitor performance comparison"],
        ],
        col_widths=[1.3, 1.7, 3.0]
    )

    add_page_break(doc)

    # 4.6
    add_heading(doc, "4.6 Settings & Preferences", level=2)

    add_styled_table(doc,
        ["Section", "What Can Be Configured"],
        [
            ["Profile", "Business name, logo, description, contact details, WhatsApp, showroom address"],
            ["Preferences", "Auto-reply ON/OFF, notification preferences, working hours, language"],
            ["Subscription", "Current plan, usage stats, upgrade/downgrade, billing history"],
            ["Team", "Add/remove team members, role assignment (Enterprise only)"],
        ],
        col_widths=[1.5, 4.5]
    )

    add_heading(doc, "Dealer Navigation", level=2)
    add_diagram_box(doc, "BOTTOM NAVIGATION BAR (Mobile)", [
        "+----------+----------+----------+----------+----------+",
        "|Dashboard |Inventory |  Leads   |  Studio  | Settings |",
        "| (grid)   |  (car)   |(people)  |  (wand)  |  (gear)  |",
        "+----------+----------+----------+----------+----------+",
    ])
    add_diagram_box(doc, "SIDEBAR (Desktop)", [
        "+----------------------+",
        "|  [CaroBest Logo]    |",
        "|                      |",
        "|  Dashboard           |",
        "|  Inventory           |",
        "|  Leads               |",
        "|  AI Studio           |",
        "|    - Photos          |",
        "|    - Content         |",
        "|    - Reels           |",
        "|    - Creative        |",
        "|  Analytics           |",
        "|    - Performance     |",
        "|    - Reports         |",
        "|  Intelligence        |",
        "|  Marketing           |",
        "|  Settings            |",
        "+----------------------+",
    ])

    add_page_break(doc)

    # ===========================
    # 5. DATA FLOW ARCHITECTURE
    # ===========================
    add_heading(doc, "5. Data Flow Architecture", level=1)

    add_heading(doc, "Database Entity Relationships", level=2)

    add_diagram_box(doc, "CORE DATA MODEL", [
        "User (Supabase Auth)",
        "  |",
        "  +---> DealerProfile (1:1)",
        "  |       |",
        "  |       +---> Vehicle[] (1:many)",
        "  |       |       +---> VehicleImage[] (1:many)",
        "  |       |",
        "  |       +---> Lead[] (1:many)",
        "  |       |       +---> LeadMessage[] (1:many)",
        "  |       |       +---> Appointment[] (1:many)",
        "  |       |",
        "  |       +---> Subscription (1:1)",
        "  |       +---> DealerPreference (1:1)",
        "  |       +---> Notification[] (1:many)",
        "  |       +---> Activity[] (1:many)",
        "  |       +---> PlatformEvent[] (1:many)",
        "  |",
        "  +---> Wishlist[] (1:many) --> Vehicle",
        "  +---> ConciergeMemory[] (1:many)",
        "  +---> ServiceBooking[] (1:many)",
        "",
        "Standalone:",
        "  CarBrand --> CarModel[] --> CarVariant[]",
        "  PlatformEvent (append-only audit log)",
    ])

    add_heading(doc, "Complete API Map", level=2)

    add_styled_table(doc,
        ["Category", "Endpoints", "Auth Required"],
        [
            ["Auth", "POST /api/auth/dealer-signup, GET /api/auth/me", "No / Yes"],
            ["Vehicles", "GET/POST /api/vehicles, GET/PATCH /api/vehicles/[id], PATCH panorama", "Varies"],
            ["Leads", "GET/POST /api/leads, GET/PATCH /api/leads/[id], messages CRUD", "Dealer"],
            ["Dealer", "GET/PATCH /api/dealer/profile, preferences", "Dealer"],
            ["AI (Text)", "description, sentiment, valuation, smart-reply, quick-draft, creative, reel", "User"],
            ["AI (Photo)", "background-remove, apply-mood", "User"],
            ["Analytics", "performance, reports, benchmarks, health-score", "Dealer"],
            ["Intelligence", "market-feed", "Dealer"],
            ["Payments", "create-order, verify", "User"],
            ["Wishlist", "GET/POST /api/wishlist", "User"],
            ["Notifications", "GET/PATCH /api/notifications", "User"],
            ["Admin", "overview, alerts, vehicles moderation", "Admin"],
            ["Cron", "nightly-scoring (Vercel Cron, 2 AM IST)", "Internal"],
        ],
        col_widths=[1.2, 3.0, 1.8]
    )

    add_page_break(doc)

    # ===========================
    # 6. AI PIPELINE
    # ===========================
    add_heading(doc, "6. AI Pipeline", level=1)

    add_heading(doc, "Centralized AI Request Flow", level=2)
    add_diagram_box(doc, "AI REQUEST LIFECYCLE", [
        "Any AI API Route calls aiRequest()",
        "          |",
        "          v",
        "  +-------+--------+",
        "  | COMPLEXITY      |",
        "  | DETECTION       |",
        "  |                 |",
        "  | SIMPLE -------> gpt-4o-mini (fast, cheap)",
        "  | MODERATE -----> gpt-4o-mini",
        "  | COMPLEX ------> gpt-4o (powerful, accurate)",
        "  +---------+------+",
        "            |",
        "            v",
        "  +-------------------+",
        "  | RETRY WITH        |",
        "  | EXPONENTIAL       |",
        "  | BACKOFF           |",
        "  |                   |",
        "  | Attempt 1: 0ms    |",
        "  | Attempt 2: ~500ms |",
        "  | Attempt 3: ~2s    |",
        "  | (30% jitter)      |",
        "  +---------+---------+",
        "            |",
        "            v",
        "  Returns: { content, model, tokensUsed, latencyMs }",
    ])

    add_heading(doc, "Circuit Breaker Protection", level=2)
    add_diagram_box(doc, "CIRCUIT BREAKER STATE MACHINE", [
        "  +----------+    3 failures    +----------+",
        "  |          | ===============> |          |",
        "  |  CLOSED  |    in 5 min      |   OPEN   |",
        "  | (normal) |                  | (failing)|",
        "  |          | <=============== |          |",
        "  +----+-----+    on success    +----+-----+",
        "       ^                             |",
        "       |                        after 30s",
        "       |                             |",
        "       |                             v",
        "       |                      +-----------+",
        "       +--- on success ------+| HALF-OPEN |",
        "                              | (probing) |",
        "       OPEN on failure <------+-----------+",
        "",
        "When OPEN:",
        "  - AI calls skip immediately (no waiting)",
        "  - Fallback response returned instantly",
        "  - Admin notified via Notification + PlatformEvent",
        "",
        "Separate circuits for: 'openai' and 'replicate'",
    ])

    add_heading(doc, "AI Features Summary", level=2)

    add_styled_table(doc,
        ["Feature", "Model", "Complexity", "Fallback When AI Unavailable"],
        [
            ["Vehicle Description", "GPT-4o-mini", "MODERATE", "Template-based text"],
            ["Lead Sentiment", "GPT-4o-mini", "MODERATE", "Message count heuristic"],
            ["Vehicle Valuation", "GPT-4o", "COMPLEX", "Deterministic price formula"],
            ["Smart Reply", "GPT-4o", "COMPLEX", "3 canned reply suggestions"],
            ["Quick Draft", "GPT-4o-mini", "MODERATE", "Template from intent"],
            ["Notification Enhance", "GPT-4o-mini", "MODERATE", "Return original text"],
            ["Creative Suggestions", "GPT-4o-mini", "MODERATE", "3 static campaign ideas"],
            ["Reel Script", "GPT-4o-mini", "MODERATE", "Template video script"],
            ["Text-to-Speech", "OpenAI TTS", "N/A", "503 Service Unavailable"],
            ["Background Remove", "Replicate rembg", "N/A", "Return original image"],
            ["Apply Mood", "Replicate SD", "N/A", "Return original image"],
            ["Concierge Chat", "GPT-4o", "COMPLEX", "Apology message"],
        ],
        col_widths=[1.3, 1.2, 1.0, 2.5]
    )

    add_page_break(doc)

    # ===========================
    # 7. EVENT SYSTEM
    # ===========================
    add_heading(doc, "7. Event System & Agent Infrastructure", level=1)

    add_heading(doc, "Event-Driven Architecture", level=2)
    add_body(doc, "Every significant action emits a PlatformEvent. Events power: audit trail, AI agent triggers, analytics, and benchmarks.")

    add_diagram_box(doc, "EVENT EMISSION FLOW", [
        "  API Route completes business logic",
        "          |",
        "          v",
        "  emitEvent({",
        "    type: 'LEAD_CREATED',",
        "    entityType: 'Lead',",
        "    entityId: lead.id,",
        "    dealerProfileId: dealer.id,",
        "    metadata: { source, buyerName }",
        "  })",
        "          |",
        "    +-----+-----+",
        "    |             |",
        "    v             v",
        "  PERSIST       PROCESS",
        "  to DB         by Agent",
        "  (async)       (async)",
        "    |             |",
        "    v             v",
        "  Audit         Trigger",
        "  Trail         AI Actions",
    ])

    add_heading(doc, "16 Event Types", level=2)

    add_styled_table(doc,
        ["Event", "Trigger", "Agent Action"],
        [
            ["LEAD_CREATED", "Buyer contacts dealer", "Auto-sentiment + auto-reply + schedule follow-ups"],
            ["LEAD_STATUS_CHANGED", "Dealer updates lead status", "Record response in agent memory"],
            ["LEAD_MESSAGE_SENT", "New message in conversation", "Track communication patterns"],
            ["SENTIMENT_ANALYZED", "AI completes sentiment analysis", "If HOT: urgent dealer notification"],
            ["DESCRIPTION_GENERATED", "AI writes vehicle description", "Track AI usage metrics"],
            ["VEHICLE_CREATED", "Dealer adds new vehicle", "Track inventory growth"],
            ["VEHICLE_WISHLISTED", "Buyer wishlists a vehicle", "If 3+ wishlists: mark as Trending"],
            ["VEHICLE_MODERATED", "Admin moderates listing", "Admin audit trail"],
            ["SUBSCRIPTION_ACTIVATED", "Payment verified", "Welcome sequence trigger"],
            ["AUTO_REPLY_SENT", "AI sends auto-reply", "Thompson Sampling memory update"],
            ["VEHICLE_SCORED", "AI scores vehicle quality", "Track scoring patterns"],
            ["TRENDING_BADGE_SET", "Vehicle marked trending", "Badge assignment audit"],
            ["CIRCUIT_BREAKER_OPEN", "AI service failing", "Admin alert notification"],
            ["API_ERROR", "Any API error occurs", "Error monitoring and alerting"],
            ["FOLLOW_UP_SENT", "Follow-up reminder sent", "Follow-up effectiveness tracking"],
            ["PRICE_ADJUSTED", "Price suggestion applied", "Price change audit trail"],
        ],
        col_widths=[1.5, 1.8, 2.7]
    )

    add_heading(doc, "Nightly Cron Job (2 AM IST daily)", level=2)

    add_styled_table(doc,
        ["#", "Task", "What It Does"],
        [
            ["1", "Re-score Vehicles", "Re-calculate AI quality scores older than 7 days"],
            ["2", "Re-analyze Leads", "Re-run sentiment on leads with no update in 48 hours"],
            ["3", "Follow-up Reminders", "Send overdue follow-up notifications to dealers"],
            ["4", "Pricing Suggestions", "Auto-adjust price recommendations based on demand"],
            ["5", "Daily Snapshot", "Take analytics snapshot for historical tracking"],
            ["6", "Activity Logging", "Record all cron actions as Activity records"],
        ],
        col_widths=[0.3, 1.7, 4.0]
    )

    add_page_break(doc)

    # ===========================
    # 8. PAYMENT FLOW
    # ===========================
    add_heading(doc, "8. Payment Flow", level=1)

    add_flow_diagram(doc, "Razorpay Subscription Payment",
        ["Select\nPlan", "Create\nOrder", "Razorpay\nCheckout", "Payment\nProcessed", "Verify\nSignature", "Plan\nActivated"])

    add_diagram_box(doc, "PAYMENT LIFECYCLE (DETAILED)", [
        "1. Dealer selects plan on /plans page",
        "2. Frontend calls POST /api/payments/create-order",
        "3. Server creates Razorpay order with plan amount",
        "4. Razorpay checkout modal opens in browser",
        "5. Dealer enters card / UPI / netbanking details",
        "6. Razorpay processes payment",
        "7. On success: frontend calls POST /api/payments/verify",
        "8. Server verifies payment signature (HMAC SHA256)",
        "9. Server creates/updates Subscription record in DB",
        "10. Server updates DealerProfile.plan field",
        "11. PlatformEvent emitted: SUBSCRIPTION_ACTIVATED",
        "12. Dealer redirected to dashboard (plan now active)",
        "",
        "DEMO MODE:",
        "If RAZORPAY_KEY_ID and RAZORPAY_KEY_SECRET are not set,",
        "the payment flow runs in demo mode (simulated success).",
        "This allows testing without real payment processing.",
    ])

    add_page_break(doc)

    # ===========================
    # 9. ERROR HANDLING
    # ===========================
    add_heading(doc, "9. Error Handling Architecture", level=1)

    add_diagram_box(doc, "ERROR HANDLING FLOW", [
        "  API ROUTE ERROR",
        "        |",
        "        v",
        "  ApiError class (typed errors)",
        "        |",
        "  +-----+----------------------------------+",
        "  |  Code              | HTTP Status        |",
        "  |  VALIDATION_ERROR  | 400 Bad Request    |",
        "  |  AUTH_REQUIRED     | 401 Unauthorized   |",
        "  |  FORBIDDEN         | 403 Forbidden      |",
        "  |  NOT_FOUND         | 404 Not Found      |",
        "  |  DB_ERROR          | 500 Server Error   |",
        "  |  AI_UNAVAILABLE    | 503 Service Unavail|",
        "  +-----+----------------------------------+",
        "        |",
        "        v",
        "  handleApiError() --> Consistent JSON response",
        "        |",
        "        +---> Sentry.captureException() (server)",
        "        +---> PlatformEvent: API_ERROR (monitoring)",
        "",
        "  CLIENT SIDE:",
        "  useApi hook catches errors",
        "        |",
        "        +---> 401 --> 'Session expired. Log in again.'",
        "        +---> 403 --> 'Access denied.'",
        "        +---> 500 --> Auto-retry (1 retry, 2s delay)",
        "        +---> Network --> 'You appear to be offline.'",
    ])

    add_page_break(doc)

    # ===========================
    # APPENDIX: AUTH & NOTIFICATIONS
    # ===========================
    add_heading(doc, "Appendix A: Authentication Flow", level=1)

    add_styled_table(doc,
        ["User Type", "Login Route", "Auth Method", "Protected Routes"],
        [
            ["Buyer", "/login/buyer", "Magic Link or Google OAuth", "/my-account/*, /wishlist, /concierge"],
            ["Dealer", "/login/dealer", "Email + Password", "/dashboard/* (all dealer pages)"],
            ["Admin", "/admin/*", "Email + Password + Admin flag", "/admin/* (all admin pages)"],
        ],
        col_widths=[1.0, 1.2, 2.0, 2.0]
    )

    add_diagram_box(doc, "SESSION MANAGEMENT", [
        "1. Supabase middleware runs on every request",
        "2. Refreshes session token automatically",
        "3. Protected routes checked by middleware",
        "4. API routes verify via supabase.auth.getUser()",
        "5. AuthGuard component wraps dealer pages (client-side)",
        "6. Redirect to login page if session missing/expired",
    ])

    add_heading(doc, "Appendix B: Notification System", level=1)

    add_styled_table(doc,
        ["Type", "Example", "Trigger"],
        [
            ["LEAD", "New lead from Rahul interested in Creta", "Buyer contacts dealer"],
            ["SYSTEM", "Your subscription renews in 3 days", "Subscription nearing expiry"],
            ["AI", "AI auto-replied to 5 leads while you were away", "Auto-reply agent activity"],
            ["ALERT", "AI service temporarily unavailable", "Circuit breaker opens"],
        ],
        col_widths=[0.8, 2.7, 2.5]
    )

    add_body(doc, "Delivery: In-app bell icon with unread count badge. Polling on dashboard every 30 seconds.")

    add_page_break(doc)

    # ===========================
    # APPENDIX: SCREEN REFERENCE
    # ===========================
    add_heading(doc, "Appendix C: Screen Design Reference", level=1)
    add_body(doc, "All UI designs are available in the stitch/ directory of the codebase. Each subfolder contains code.html (implementation) and screen.png (visual reference).")

    add_heading(doc, "Buyer Screens", level=2)
    add_styled_table(doc,
        ["Screen", "Stitch Folder", "Description"],
        [
            ["Landing Page", "stitch/homepage/", "Hero + featured cars + trust signals"],
            ["Showroom", "stitch/showroom/", "Vehicle grid with filters and sort"],
            ["Vehicle Detail", "stitch/vehicle-detail/", "Full car detail page"],
            ["Compare", "stitch/compare/", "Side-by-side vehicle comparison"],
            ["Wishlist", "stitch/wishlist/", "Saved vehicles list"],
            ["Concierge", "stitch/concierge/", "AI chat interface"],
            ["My Account", "stitch/my-account/", "Account sections hub"],
            ["Car Loan", "stitch/car-loan/", "EMI calculator"],
            ["Swap", "stitch/swap/", "Trade-in valuation flow"],
        ],
        col_widths=[1.2, 1.8, 3.0]
    )

    add_heading(doc, "Dealer Screens", level=2)
    add_styled_table(doc,
        ["Screen", "Stitch Folder", "Description"],
        [
            ["Dashboard", "stitch/dashboard/", "Metrics + activity feed"],
            ["Inventory", "stitch/inventory/", "Vehicle list + management"],
            ["Leads", "stitch/leads/", "Lead list + detail view"],
            ["AI Studio", "stitch/studio/", "Photo + content AI tools"],
            ["Analytics", "stitch/analytics/", "Performance charts + reports"],
            ["Intelligence", "stitch/intelligence/", "Market data + insights"],
            ["Settings", "stitch/settings/", "Profile + preferences"],
            ["Plans", "stitch/plans/", "Subscription plan selection"],
        ],
        col_widths=[1.2, 1.8, 3.0]
    )

    # ===========================
    # SAVE
    # ===========================
    output_path = "/Users/carobest-web/public/carobest-user-journey.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print(f"Pages: ~30+ (estimated)")

generate()

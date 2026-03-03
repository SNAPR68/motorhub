#!/usr/bin/env python3
"""Generate CaroBest Platform Guide as a professional Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

BLUE = RGBColor(0x11, 0x52, 0xD4)
DARK = RGBColor(0x10, 0x16, 0x22)
GOLD = RGBColor(0xD4, 0xA0, 0x17)
GREEN = RGBColor(0x10, 0xB9, 0x81)
GRAY = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF1, 0xF5, 0xF9)

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

for level in range(1, 5):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = DARK

doc.styles["Heading 1"].font.size = Pt(22)
doc.styles["Heading 1"].font.bold = True
doc.styles["Heading 2"].font.size = Pt(16)
doc.styles["Heading 2"].font.bold = True
doc.styles["Heading 3"].font.size = Pt(13)
doc.styles["Heading 3"].font.bold = True


def add_colored_heading(text, level=1, color=BLUE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def add_body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def set_cell_bg(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    s = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(s)


def add_table(headers, rows, col_widths=None, header_color="1152D4"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, header_color)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_bg(cell, "F8FAFC")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)

    return tbl


def add_divider():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    run.font.size = Pt(8)


# ═══════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════

doc.add_paragraph()  # spacer
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CAROBEST")
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Platform Guide")
run.font.size = Pt(24)
run.font.color.rgb = DARK

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Complete Dealer & Buyer Experience Documentation")
run.font.size = Pt(13)
run.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Version 2.0  |  March 2026")
run.font.size = Pt(11)
run.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI-Powered Used Car Marketplace for India")
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = GRAY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════

add_colored_heading("Table of Contents", level=1, color=DARK)

toc_items = [
    ("1.", "Platform Overview"),
    ("2.", "Dealer Experience"),
    ("  2.1", "Registration & Onboarding"),
    ("  2.2", "Dealer Dashboard (DealerOS)"),
    ("  2.3", "Inventory Management"),
    ("  2.4", "Lead Management & CRM"),
    ("  2.5", "AI Studio (6 Tools)"),
    ("  2.6", "Analytics & Reporting"),
    ("  2.7", "DemandPulse Intelligence"),
    ("  2.8", "Marketing & Social Hub"),
    ("  2.9", "Store & Team Management"),
    ("  2.10", "Settings & Configuration"),
    ("  2.11", "Subscription Plans & Billing"),
    ("3.", "Buyer Experience"),
    ("  3.1", "Discovery & Browsing"),
    ("  3.2", "Vehicle Detail & Decision Tools"),
    ("  3.3", "Comparison Engine"),
    ("  3.4", "Wishlist & Interests"),
    ("  3.5", "AI Concierge"),
    ("  3.6", "VIP Programme"),
    ("  3.7", "Services Marketplace"),
    ("  3.8", "Post-Purchase Hub (My Account)"),
    ("  3.9", "Finance & Insurance Tools"),
    ("4.", "How Dealer & Buyer Connect"),
    ("5.", "Subscription Plans (Full Comparison)"),
    ("6.", "Technical Architecture"),
    ("7.", "API Reference Summary"),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(f"{num}  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BLUE if not num.startswith("  ") else GRAY
    r2 = p.add_run(title)
    r2.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. PLATFORM OVERVIEW
# ═══════════════════════════════════════════════════════════

add_colored_heading("1. Platform Overview", level=1)

add_body("CaroBest is an AI-powered used car marketplace built for the Indian market. It serves two distinct user types through a unified platform:")

add_bullet("operate through a dedicated portal (DealerOS) to manage inventory, leads, AI-generated content, analytics, and marketing. Access is restricted behind authentication and route protection.", bold_prefix="Dealers ")
add_bullet("browse vehicles, compare options, use AI-powered decision tools, access services (financing, insurance, RC transfer), and manage post-purchase needs.", bold_prefix="Buyers ")

add_body("The two experiences share a common vehicle database: dealers list vehicles through DealerOS, and those vehicles surface in the buyer marketplace. Buyer inquiries generate leads that flow into the dealer CRM.")

add_colored_heading("Platform at a Glance", level=3, color=DARK)

add_table(
    ["Metric", "Value"],
    [
        ["Total Routes", "204+"],
        ["API Endpoints", "37+"],
        ["Screen Designs Implemented", "75+"],
        ["Killer Features", "10"],
        ["Subscription Tiers", "4 (FREE / STARTER / GROWTH / ENTERPRISE)"],
        ["Database Models", "12 core models"],
        ["AI Integrations", "4 (Descriptions, Concierge, Valuation, Sentiment)"],
    ],
    col_widths=[2.5, 4.0],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 2. DEALER EXPERIENCE
# ═══════════════════════════════════════════════════════════

add_colored_heading("2. Dealer Experience", level=1)

# 2.1 Registration & Onboarding
add_colored_heading("2.1 Registration & Onboarding", level=2)

add_body("Entry Points:", bold=True)
add_bullet("Signup: /dealer/signup")
add_bullet("Login: /login/dealer")
add_bullet("Google SSO: Available on both pages")

add_colored_heading("Signup Form", level=3, color=DARK)
add_body("The dealer signup page collects all information needed to create a dealership account:")

add_table(
    ["Section", "Fields"],
    [
        ["Business Details", "Dealership name, GST number, Business type (Individual / Partnership / Pvt Ltd / LLP), City, State"],
        ["Owner Details", "Full name, Email address, Phone number"],
        ["Security", "Password, Confirm password"],
        ["Plan Selection", "FREE, Starter (INR 1,999/mo), Growth (INR 4,999/mo), Enterprise (INR 14,999/mo)"],
    ],
    col_widths=[1.8, 4.7],
)

add_body("")
add_body("After successful submission, the system creates a Supabase auth user + DealerProfile record via the /api/auth/dealer-signup endpoint (Zod-validated), then redirects to the dealer login page with an onboarding flag.")

add_colored_heading("5-Step Onboarding Wizard", level=3, color=DARK)
add_body("After first login, new dealers are guided through a comprehensive onboarding flow at /onboarding/dealer:")

add_table(
    ["Step", "Name", "What Happens"],
    [
        ["1", "Store Setup", "Configure store display name, upload logo, set operating hours, add WhatsApp contact number"],
        ["2", "Add Inventory", "Add first car via photo upload (AI auto-fills details) or manual entry; option to skip"],
        ["3", "Connect Social", "Link Instagram (auto-post), Facebook (Marketplace), WhatsApp Business (auto-reply), YouTube (walkarounds)"],
        ["4", "AI Tools", "Configure AI description style, enable auto-reply, set smart pricing, activate lead scoring"],
        ["5", "Go Live", "Final review and publish storefront to the marketplace"],
    ],
    col_widths=[0.5, 1.2, 4.8],
)

add_body("")
add_body("After completing onboarding, the dealer lands on their executive dashboard.")

add_divider()

# 2.2 Dashboard
add_colored_heading("2.2 Dealer Dashboard (DealerOS)", level=2)
add_body("Route: /dashboard", italic=True)
add_body("")
add_body("The executive dashboard provides a command center for dealership operations with real-time data from the API.")

add_colored_heading("Dashboard Metrics", level=3, color=DARK)
add_table(
    ["Metric", "Source", "Description"],
    [
        ["Portfolio Value", "API: /api/dashboard", "Total revenue from all vehicle sales"],
        ["Active Leads", "API: /api/dashboard", "Number of leads in active pipeline"],
        ["Hot Leads", "API: /api/dashboard", "AI-scored high-intent leads"],
        ["AI Replies Sent", "API: /api/dashboard", "Auto-replies generated by AI"],
        ["Total Vehicles", "API: /api/dashboard", "Vehicles currently in inventory"],
    ],
    col_widths=[1.5, 1.8, 3.2],
)

add_colored_heading("Quick Actions", level=3, color=DARK)
add_table(
    ["Action", "Route", "Description"],
    [
        ["Reports", "/reports/monthly", "Monthly performance reports"],
        ["Messages", "/leads", "Lead inbox / CRM"],
        ["Schedule", "/appointments", "Appointment calendar"],
        ["Market", "/marketing", "Marketing tools"],
        ["Intelligence", "/intelligence", "DemandPulse market intelligence"],
        ["Analytics", "/analytics", "Performance analytics"],
        ["Social Hub", "/social-hub", "Social media management"],
    ],
    col_widths=[1.3, 2.0, 3.2],
)

add_body("")
add_body("Additional dashboard variants:", bold=True)
add_bullet("/dashboard/overview -- Expanded view with performance rings, appointment timeline, and activity feed")
add_bullet("/dashboard/ai-hub -- AI command center with quick links to all AI tools")

add_divider()

# 2.3 Inventory
add_colored_heading("2.3 Inventory Management", level=2)
add_body("Routes: /inventory, /inventory/add, /inventory/[id]", italic=True)
add_body("")

add_body("Full vehicle lifecycle management with AI-powered enhancements:")

add_bullet("Vehicle CRUD operations (Create, Read, Update, Delete)")
add_bullet("Status management: AVAILABLE / SOLD / RESERVED")
add_bullet("AI-powered vehicle descriptions auto-generated from photos and specs")
add_bullet("Photo management with AI studio integration")
add_bullet("AI Score (0-100) assessing listing quality")
add_bullet("Price display formatting (INR with commas)")
add_bullet("Bulk operations for multi-vehicle actions")
add_bullet("Filter by status, price range, make/model; sort by date, price, AI score")

add_colored_heading("Vehicle Data Model", level=3, color=DARK)
add_body("Each vehicle record includes: make, model, year, variant, fuel type, transmission, kilometers driven, ownership count, registration state, color, price, AI-generated description, up to 20 photos, VIN, registration number, insurance validity, and AI quality score.")

add_divider()

# 2.4 Lead Management
add_colored_heading("2.4 Lead Management & CRM", level=2)
add_body("Routes: /leads, /leads/[id]", italic=True)
add_body("")

add_colored_heading("Lead Pipeline", level=3, color=DARK)
add_table(
    ["Status", "Description"],
    [
        ["NEW", "Fresh inquiry, not yet contacted"],
        ["CONTACTED", "Initial outreach made"],
        ["QUALIFIED", "Buyer shows genuine intent"],
        ["CLOSED_WON", "Sale completed"],
        ["CLOSED_LOST", "Lead did not convert"],
    ],
    col_widths=[1.8, 4.7],
)

add_colored_heading("Lead Intelligence", level=3, color=DARK)
add_table(
    ["Feature", "Description", "Min Plan"],
    [
        ["Sentiment Analysis", "AI scores lead intent as HOT / WARM / COOL", "STARTER"],
        ["Source Tracking", "WEBSITE / FACEBOOK / INSTAGRAM / WHATSAPP / WALKIN / REFERRAL", "FREE"],
        ["Conversation Thread", "Full message history (manual + AI auto-replies)", "FREE"],
        ["Smart Reply", "AI-generated contextual response suggestions", "GROWTH"],
        ["Auto-Reply", "Instant AI response to new inquiries", "GROWTH"],
        ["Follow-up Scheduling", "AI-timed follow-up reminders", "GROWTH"],
        ["Agent Memory", "Personalized interactions across conversations", "GROWTH"],
        ["Lead Assignment", "Assign leads to specific team members", "GROWTH"],
    ],
    col_widths=[1.5, 3.3, 1.5],
)

add_body("")
add_body("Additional tools:", bold=True)
add_bullet("Smart Reply (/smart-reply) -- AI-powered response suggestions analyzing lead context and vehicle interest")
add_bullet("Quick Draft (/quick-draft) -- AI content drafting for messages, descriptions, and marketing copy")

add_divider()

# 2.5 AI Studio
add_colored_heading("2.5 AI Studio", level=2)
add_body("Route: /studio", italic=True)
add_body("")
add_body("The AI Studio is the creative hub providing six specialized tools:")

add_table(
    ["Tool", "Description", "Min Plan"],
    [
        ["AI Description Generator", "Auto-generate compelling vehicle listings from photos and specs", "FREE"],
        ["Background Removal", "Remove/replace photo backgrounds with studio-quality results", "STARTER"],
        ["Mood Application", "Apply cinematic lighting and mood presets to vehicle photos", "GROWTH"],
        ["Reel Script Generator", "Create video scripts for social media reels with hooks and CTAs", "GROWTH"],
        ["Text-to-Speech", "Convert scripts to voiceover audio for video content", "GROWTH"],
        ["Cinema Exports", "4K export with cinematic color grading", "ENTERPRISE"],
    ],
    col_widths=[1.8, 3.5, 1.2],
)

add_body("")
add_body("Extended creative tools:", bold=True)
add_bullet("Content Studio (/content-studio) -- Marketing materials workspace for social posts, email templates, and promotional content")
add_bullet("Reel Editor (/reel-editor) -- Video creation tool with script-to-reel workflow, music library, transition effects")

add_divider()

# 2.6 Analytics
add_colored_heading("2.6 Analytics & Reporting", level=2)

add_table(
    ["Page", "Route", "Description"],
    [
        ["Analytics Dashboard", "/analytics", "Enquiry trends, conversion funnel, source performance, vehicle metrics"],
        ["Monthly Reports", "/reports/monthly", "MoM growth, revenue tracking, lead source breakdown, top vehicles"],
        ["Performance", "/performance", "Inventory health scores, days on lot, price competitiveness"],
        ["Profit Tracking", "/profit", "Purchase vs. sale prices, per-vehicle margins, monthly trends"],
        ["Report Hub", "/reports", "Central access to all report types"],
    ],
    col_widths=[1.5, 1.8, 3.2],
)

add_body("")
add_body("Analytics retention by plan:", bold=True)
add_bullet("FREE: 7 days")
add_bullet("STARTER: 30 days")
add_bullet("GROWTH: 90 days + Advanced Analytics + Benchmarks + Health Score")
add_bullet("ENTERPRISE: Unlimited retention + Data Export (CSV/PDF)")

add_divider()

# 2.7 Intelligence
add_colored_heading("2.7 DemandPulse Intelligence", level=2)
add_body("Route: /intelligence", italic=True)
add_body("")
add_body("Market intelligence suite providing data-driven insights for pricing and inventory decisions.")

add_table(
    ["Page", "Route", "Description", "Min Plan"],
    [
        ["Overview", "/intelligence", "Market summary with key indicators", "GROWTH"],
        ["Pricing Intelligence", "/intelligence/pricing", "Real-time pricing data, competitor analysis, recommendations", "GROWTH"],
        ["Demand Trends", "/intelligence/trends", "Search volume trends, popular models, emerging demand", "GROWTH"],
        ["Depreciation Curves", "/intelligence/depreciation", "Model-wise depreciation, optimal buy/sell timing", "GROWTH"],
        ["Market Forecast", "/intelligence/forecast", "Predictive analytics for market movements", "ENTERPRISE"],
        ["Competitor Analysis", "/intelligence/competitors", "Competitor inventory monitoring, pricing comparison", "ENTERPRISE"],
    ],
    col_widths=[1.2, 1.8, 2.5, 1.0],
)

add_divider()

# 2.8 Marketing
add_colored_heading("2.8 Marketing & Social Hub", level=2)

add_table(
    ["Feature", "Route", "Description", "Min Plan"],
    [
        ["Marketing Hub", "/marketing", "Campaign creation, performance tracking, budget allocation", "GROWTH"],
        ["Campaigns", "/marketing/campaigns", "Campaign management and scheduling", "GROWTH"],
        ["Templates", "/marketing/templates", "Marketing template library", "GROWTH"],
        ["Social Hub", "/social-hub", "Unified social media management, content calendar, analytics", "GROWTH"],
        ["Auto-Posting", "(within Social Hub)", "Scheduled automatic posting to linked platforms", "ENTERPRISE"],
    ],
    col_widths=[1.2, 1.8, 2.5, 1.0],
)

add_divider()

# 2.9 Store & Team
add_colored_heading("2.9 Store & Team Management", level=2)

add_colored_heading("Store Limits", level=3, color=DARK)
add_table(
    ["Plan", "Stores", "Team Members", "RBAC"],
    [
        ["FREE", "1", "1", "No"],
        ["STARTER", "1", "2", "No"],
        ["GROWTH", "3", "5", "Yes"],
        ["ENTERPRISE", "Unlimited", "Unlimited", "Yes"],
    ],
    col_widths=[1.5, 1.5, 1.5, 1.5],
)

add_body("")
add_body("Multi-Store Management (/stores):", bold=True)
add_bullet("Add and manage multiple dealership locations")
add_bullet("Per-store inventory tracking and analytics")

add_body("Team Management (/settings/team):", bold=True)
add_bullet("Invite by email with role assignment: Owner / Manager / Salesperson / Viewer")
add_bullet("Per-member lead assignment and activity tracking")

add_body("Appointments (/appointments):", bold=True)
add_bullet("Calendar-based scheduling with customer name, vehicle interest, date/time, notes")
add_bullet("Status: Scheduled / Completed / Cancelled / No-show")

add_divider()

# 2.10 Settings
add_colored_heading("2.10 Settings & Configuration", level=2)
add_body("Route: /settings", italic=True)

add_table(
    ["Setting", "Route", "Description"],
    [
        ["Profile", "/settings", "Dealership info, logo, contact details"],
        ["Team", "/settings/team", "Team member management + RBAC"],
        ["Billing", "/settings/billing", "Subscription management, invoices"],
        ["Permissions", "/settings/permissions", "Feature access matrix based on current plan"],
        ["AI Permissions", "/settings/ai-permissions", "Granular AI feature toggles"],
        ["Automation", "/settings/automation", "Auto-reply rules, follow-up schedules"],
        ["Brand Voice", "/settings/brand-voice", "AI tone/style configuration"],
        ["Notifications", "/settings/notifications", "Notification delivery preferences"],
        ["Assets", "/settings/assets", "Brand assets (logos, watermarks)"],
    ],
    col_widths=[1.5, 2.2, 2.8],
)

add_divider()

# 2.11 Plans
add_colored_heading("2.11 Subscription Plans & Billing", level=2)
add_body("Routes: /plans (comparison), /settings/billing (current subscription)", italic=True)
add_body("")
add_body("Full plan comparison is covered in Section 5.")
add_body("Payment processing: Razorpay integration with order creation, payment verification, and demo mode fallback.")
add_body("Annual billing: 2 months free on all paid plans.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 3. BUYER EXPERIENCE
# ═══════════════════════════════════════════════════════════

add_colored_heading("3. Buyer Experience", level=1)

# 3.1 Discovery
add_colored_heading("3.1 Discovery & Browsing", level=2)

add_table(
    ["Entry Point", "Route", "Description"],
    [
        ["Homepage", "/", "Featured vehicles, trending cars, brand showcase, city-wise browsing"],
        ["Landing", "/landing", "Cinematic hero with Browse Collection and Private Showroom CTAs"],
        ["Showroom", "/showroom", "Curated vehicle gallery with premium presentation"],
        ["Used Cars", "/used-cars", "Full marketplace with advanced filters"],
        ["New Cars", "/new-cars", "New car catalog browsing"],
    ],
    col_widths=[1.3, 1.5, 3.7],
)

add_body("")
add_body("Homepage features:", bold=True)
add_bullet("Hero section with search functionality")
add_bullet("Trending vehicles carousel")
add_bullet("Brand grid: 12 brands (Maruti, Hyundai, Tata, Honda, Toyota, Mahindra, Kia, MG, Volkswagen, Skoda, BMW, Mercedes)")
add_bullet("City-wise browsing, featured dealers, quick-access services")

add_body("Filter capabilities:", bold=True)
add_bullet("Price range slider, multi-select brand/fuel/transmission/body type")
add_bullet("Location-based filtering, sort by relevance/price/year/mileage")
add_bullet("Grid and list view toggles")

add_body("SEO Pages:", bold=True)
add_bullet("/used-cars/[city] -- City-specific listings")
add_bullet("/[brand]/[model] -- Brand and model pages")
add_bullet("/dealers, /dealers/[city] -- Dealer directory")
add_bullet("/car-news -- Automotive news and reviews")

add_divider()

# 3.2 Vehicle Detail
add_colored_heading("3.2 Vehicle Detail & Decision Tools", level=2)
add_body("Route: /vehicle/[id]", italic=True)
add_body("")

add_body("Vehicle detail page includes full-screen gallery, price with EMI estimate, key specs grid, AI description, dealer info card, similar cars, share and wishlist actions.")

add_colored_heading("3 Exclusive Decision Tools", level=3, color=DARK)

add_table(
    ["Tool", "Route", "Capabilities"],
    [
        ["TrueCost Engine", "/vehicle/[id]/true-cost", "Purchase price + registration + insurance + maintenance + fuel + depreciation = 5-year TCO"],
        ["AI Negotiation Coach", "/vehicle/[id]/negotiate", "Fair market price, talking points, walk-away price, counter-offer strategies"],
        ["Vehicle Passport", "/vehicle/passport/[id]", "Ownership history, service records, insurance claims, RTO verification, condition score"],
    ],
    col_widths=[1.5, 2.0, 3.0],
)

add_divider()

# 3.3 Comparison
add_colored_heading("3.3 Comparison Engine", level=2)

add_table(
    ["View", "Route", "Description"],
    [
        ["Comparison Hub", "/compare", "Select and compare up to 3 vehicles"],
        ["Technical Specs", "/compare/technical", "Side-by-side specifications table"],
        ["Visual Compare", "/compare/studio", "Photo side-by-side with overlay comparison"],
    ],
    col_widths=[1.5, 2.0, 3.0],
)

add_divider()

# 3.4 Wishlist
add_colored_heading("3.4 Wishlist & Interests", level=2)

add_bullet("Wishlist (/wishlist) -- Save vehicles, price drop alerts, quick comparison")
add_bullet("Interests (/interests) -- AI-powered recommendations based on browsing history")

add_divider()

# 3.5 AI Concierge
add_colored_heading("3.5 AI Concierge", level=2)

add_table(
    ["Mode", "Route", "Description"],
    [
        ["Chat Interface", "/concierge", "Conversational AI for vehicle discovery and questions"],
        ["Discovery Mode", "/concierge/discovery", "Guided questionnaire: budget, use case, fuel preference, must-haves; returns AI shortlist"],
        ["Vehicle Dossier", "/concierge/dossier", "AI-compiled report with market analysis, cost projection, buy/wait recommendation"],
    ],
    col_widths=[1.3, 2.0, 3.2],
)

add_divider()

# 3.6 VIP
add_colored_heading("3.6 VIP Programme", level=2)

add_table(
    ["Page", "Route", "Description"],
    [
        ["VIP Overview", "/vip", "Membership details and benefits"],
        ["VIP Showroom", "/vip/showroom", "Exclusive VIP-only vehicle showcase"],
        ["Rewards", "/vip/rewards", "Points tracking and reward redemption"],
        ["Confirmation", "/vip/confirmation", "VIP enrollment confirmation"],
    ],
    col_widths=[1.3, 2.0, 3.2],
)

add_body("")
add_body("Benefits: priority access, exclusive inventory, enhanced negotiation tools, concierge priority, reward points, special financing rates.")

add_divider()

# 3.7 Services
add_colored_heading("3.7 Services Marketplace", level=2)

add_table(
    ["Service", "Route", "Description"],
    [
        ["SwapDirect", "/swap, /swap/deal, /swap/matches", "Direct car exchange marketplace with AI-matched opportunities"],
        ["LiveCondition", "/inspection", "200+ checkpoint professional vehicle inspection booking"],
        ["InstantRC", "/rc-transfer", "Online RC transfer with document upload and status tracking"],
        ["CrossState Express", "/cross-state", "Inter-state transfer with NOC, road tax calculation"],
        ["Sell Your Car", "/sell-car", "Sell-your-car flow with AI valuation"],
    ],
    col_widths=[1.5, 2.3, 2.7],
)

add_divider()

# 3.8 Post-Purchase
add_colored_heading("3.8 Post-Purchase Hub (My Account)", level=2)
add_body("Route: /my-account", italic=True)

add_table(
    ["Section", "Route", "Description"],
    [
        ["Profile", "/my-account", "Personal info, contact details, preferences"],
        ["My Garage", "/my-account/garage", "Owned vehicles, service reminders, mileage tracking"],
        ["Warranty", "/my-account/warranty", "Warranty tracking and claims"],
        ["Documents", "/my-account/documents", "Digital vault for RC, insurance, PUC with expiry alerts"],
        ["Orders", "/my-account/orders", "Purchase history and order tracking"],
        ["Resale", "/my-account/resale", "Resale value tracking and re-listing tools"],
    ],
    col_widths=[1.2, 2.2, 3.1],
)

add_divider()

# 3.9 Finance
add_colored_heading("3.9 Finance & Insurance Tools", level=2)

add_table(
    ["Tool", "Route", "Description"],
    [
        ["Car Loan Calculator", "/car-loan", "EMI calculator, bank comparison, pre-approval check"],
        ["Car Insurance", "/car-insurance", "Comprehensive vs. third-party, IDV calculator, multi-insurer quotes"],
        ["Fuel Price", "/fuel-price", "City-wise real-time fuel price tracker"],
        ["RTO Info", "/rto", "RTO office directory and fee calculator"],
    ],
    col_widths=[1.5, 1.5, 3.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 4. HOW DEALER & BUYER CONNECT
# ═══════════════════════════════════════════════════════════

add_colored_heading("4. How Dealer & Buyer Connect", level=1)

add_body("The platform uses a shared PostgreSQL database (Supabase) as the connection point between dealer and buyer experiences.")

add_colored_heading("Vehicle Listing Flow", level=3, color=DARK)

flow_steps = [
    ("1.", "Dealer adds vehicle via /inventory/add"),
    ("2.", "Vehicle stored in database (Vehicle table, status: AVAILABLE, linked to DealerProfile)"),
    ("3.", "AI auto-generates vehicle description"),
    ("4.", "Buyer browses /showroom, /used-cars, /new-cars"),
    ("5.", "API returns AVAILABLE vehicles (GET /api/vehicles)"),
    ("6.", "Buyer views vehicle at /vehicle/[id]"),
    ("7.", "Buyer inquires (WhatsApp, call, or in-app message)"),
    ("8.", "Lead created in database (Lead table, linked to Vehicle + buyer contact)"),
    ("9.", "AI analyzes lead sentiment (HOT / WARM / COOL)"),
    ("10.", "Lead appears in dealer CRM at /leads"),
    ("11.", "Dealer responds (manual or AI auto-reply)"),
    ("12.", "Conversation tracked in LeadMessage table"),
    ("13.", "Deal progresses: NEW -> CONTACTED -> QUALIFIED -> CLOSED_WON"),
]

for num, text in flow_steps:
    p = doc.add_paragraph()
    r = p.add_run(f"{num} ")
    r.bold = True
    r.font.color.rgb = BLUE
    p.add_run(text)

add_colored_heading("Data Isolation", level=3, color=DARK)

add_bullet("Buyers never see dealer-internal pages (server middleware + AuthGuard protection)")
add_bullet("No dealer CTAs, links, or references appear in the buyer UI")
add_bullet("29 dealer route prefixes are protected by middleware -- unauthenticated access redirects to /login/dealer")
add_bullet("Buyer pages use BuyerBottomNav and BuyerPageLayout; dealer pages use DealerAppShell and DealerPageLayout")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 5. SUBSCRIPTION PLANS
# ═══════════════════════════════════════════════════════════

add_colored_heading("5. Subscription Plans (Full Comparison)", level=1)

add_colored_heading("Pricing", level=3, color=DARK)
add_table(
    ["", "FREE", "STARTER", "GROWTH", "ENTERPRISE"],
    [
        ["Monthly Price", "INR 0", "INR 1,999", "INR 4,999", "INR 14,999"],
        ["Annual Price", "INR 0", "INR 19,990", "INR 49,990", "INR 1,49,990"],
        ["Annual Savings", "--", "INR 3,998", "INR 9,998", "INR 29,998"],
        ["Tagline", "Get started", "Independent dealers", "Growing dealerships", "Large chains"],
        ["Public Badge", "--", "--", "Gold Dealer", "Platinum Dealer"],
    ],
    col_widths=[1.3, 1.2, 1.2, 1.4, 1.4],
)

add_body("")
add_colored_heading("Quotas", level=3, color=DARK)
add_table(
    ["Resource", "FREE", "STARTER", "GROWTH", "ENTERPRISE"],
    [
        ["Vehicle Listings", "5", "25", "100", "Unlimited"],
        ["Leads / month", "20", "100", "Unlimited", "Unlimited"],
        ["AI Calls / month", "10", "50", "200", "Unlimited"],
        ["Photo Edits / month", "0", "20", "100", "Unlimited"],
        ["Video Scripts / month", "0", "0", "30", "Unlimited"],
        ["Team Members", "1", "2", "5", "Unlimited"],
        ["Analytics Retention", "7 days", "30 days", "90 days", "Unlimited"],
        ["Stores", "1", "1", "3", "Unlimited"],
    ],
    col_widths=[1.5, 1.1, 1.1, 1.2, 1.2],
)

add_body("")
add_colored_heading("Features (33 flags)", level=3, color=DARK)

# Split into categories for readability
feature_sections = [
    ("AI Content", [
        ["AI Descriptions", "Yes", "Yes", "Yes", "Yes"],
        ["AI Quick Draft", "--", "Yes", "Yes", "Yes"],
        ["AI Creative Suggestions", "--", "--", "Yes", "Yes"],
        ["AI Notification Enhance", "--", "--", "Yes", "Yes"],
    ]),
    ("AI Leads", [
        ["AI Sentiment", "--", "Yes", "Yes", "Yes"],
        ["AI Smart Reply", "--", "--", "Yes", "Yes"],
        ["AI Auto-Reply", "--", "--", "Yes", "Yes"],
        ["AI Follow-ups", "--", "--", "Yes", "Yes"],
        ["Agent Memory", "--", "--", "Yes", "Yes"],
    ]),
    ("AI Media", [
        ["Background Removal", "--", "Yes", "Yes", "Yes"],
        ["Mood Application", "--", "--", "Yes", "Yes"],
        ["Reel Scripts", "--", "--", "Yes", "Yes"],
        ["Text-to-Speech", "--", "--", "Yes", "Yes"],
        ["Cinema Exports (4K)", "--", "--", "--", "Yes"],
    ]),
    ("Analytics & Intelligence", [
        ["Advanced Analytics", "--", "--", "Yes", "Yes"],
        ["Benchmarks", "--", "--", "Yes", "Yes"],
        ["Health Score", "--", "--", "Yes", "Yes"],
        ["Intelligence (Basic)", "--", "--", "Yes", "Yes"],
        ["Intelligence (Full)", "--", "--", "--", "Yes"],
        ["Data Export", "--", "--", "--", "Yes"],
    ]),
    ("Communication", [
        ["WhatsApp Integration", "--", "--", "Yes", "Yes"],
        ["Bulk Messaging", "--", "--", "--", "Yes"],
    ]),
    ("Marketing", [
        ["Social Hub", "--", "--", "Yes", "Yes"],
        ["Campaign Tools", "--", "--", "Yes", "Yes"],
        ["Auto-Posting", "--", "--", "--", "Yes"],
    ]),
    ("Management", [
        ["Lead Assignment", "--", "--", "Yes", "Yes"],
        ["Multi-Store", "--", "--", "--", "Yes"],
        ["Team Roles (RBAC)", "--", "--", "Yes", "Yes"],
        ["API Access", "--", "--", "--", "Yes"],
        ["Custom Branding", "--", "--", "--", "Yes"],
    ]),
    ("Virtual Tour & Profile", [
        ["360 Virtual Tour", "--", "Yes", "Yes", "Yes"],
        ["Interior Viewer", "--", "--", "Yes", "Yes"],
        ["Remove Branding", "--", "Yes", "Yes", "Yes"],
    ]),
]

for section_name, rows in feature_sections:
    p = doc.add_paragraph()
    r = p.add_run(section_name)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BLUE

    add_table(
        ["Feature", "FREE", "STARTER", "GROWTH", "ENTERPRISE"],
        rows,
        col_widths=[1.5, 1.1, 1.1, 1.2, 1.2],
    )
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 6. TECHNICAL ARCHITECTURE
# ═══════════════════════════════════════════════════════════

add_colored_heading("6. Technical Architecture", level=1)

add_colored_heading("Technology Stack", level=3, color=DARK)
add_table(
    ["Component", "Technology"],
    [
        ["Framework", "Next.js 16.1.6 (App Router)"],
        ["Language", "TypeScript 5"],
        ["UI Library", "React 19"],
        ["Styling", "Tailwind CSS v4"],
        ["Database", "PostgreSQL via Supabase"],
        ["ORM", "Prisma 7.4.1"],
        ["Authentication", "Supabase Auth (email/password + Google SSO)"],
        ["AI Provider", "OpenAI (descriptions, concierge, valuation, sentiment)"],
        ["Media AI", "Replicate (photo studio)"],
        ["Payments", "Razorpay (subscriptions)"],
        ["State Management", "Zustand 5"],
        ["Animations", "Framer Motion 12"],
        ["Icons", "Material Symbols Outlined"],
    ],
    col_widths=[1.8, 4.7],
)

add_colored_heading("Route Protection", level=3, color=DARK)
add_body("Server-side middleware (src/lib/supabase/middleware.ts) protects all dealer routes:")
add_bullet("Refreshes the auth session on every request")
add_bullet("Checks 29 protected route prefixes (/dashboard, /inventory, /leads, /settings, /studio, /intelligence, etc.)")
add_bullet("Redirects unauthenticated users to /login/dealer?redirect={path}")

add_colored_heading("Data Fetching Pattern", level=3, color=DARK)
add_body("Client Component -> useApi(fetchFn, deps) -> src/lib/api.ts -> /api/* routes -> Prisma ORM -> Supabase PostgreSQL")

add_colored_heading("Key Database Models", level=3, color=DARK)
add_table(
    ["Model", "Description"],
    [
        ["User", "Auth user (buyer or dealer)"],
        ["DealerProfile", "One-to-one with User; plan, dealershipName, GST, etc."],
        ["Vehicle", "Car listing with AI score, status, pricing"],
        ["Lead", "Buyer inquiry with sentiment, source, status pipeline"],
        ["LeadMessage", "Conversation messages (manual + AI auto-replies)"],
        ["Notification", "System notifications"],
        ["Activity", "Audit trail for all actions"],
        ["Appointment", "Scheduled meetings"],
        ["Subscription", "Plan billing records"],
        ["Store", "Multi-location support"],
    ],
    col_widths=[1.5, 5.0],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 7. API REFERENCE
# ═══════════════════════════════════════════════════════════

add_colored_heading("7. API Reference Summary", level=1)

api_sections = [
    ("Authentication", [
        ["GET", "/api/auth/me", "Current user + dealer profile"],
        ["POST", "/api/auth/dealer-signup", "Create dealer account (Zod-validated)"],
    ]),
    ("Vehicles", [
        ["GET", "/api/vehicles", "Paginated vehicle list (filter by status)"],
        ["GET", "/api/vehicles/[id]", "Single vehicle detail"],
        ["POST", "/api/vehicles", "Create vehicle (dealer auth)"],
        ["PUT", "/api/vehicles/[id]", "Update vehicle (dealer auth)"],
        ["DELETE", "/api/vehicles/[id]", "Delete vehicle (dealer auth)"],
    ]),
    ("Leads", [
        ["GET", "/api/leads", "Lead list (dealer auth)"],
        ["GET", "/api/leads/[id]", "Lead detail + messages"],
        ["POST", "/api/leads", "Create lead (buyer inquiry)"],
        ["PATCH", "/api/leads/[id]", "Update lead status"],
    ]),
    ("Analytics", [
        ["GET", "/api/analytics/performance", "Inventory health, sentiment, revenue"],
        ["GET", "/api/analytics/reports", "Monthly stats, MoM growth"],
        ["GET", "/api/dashboard", "Dashboard summary metrics"],
    ]),
    ("Dealer", [
        ["GET", "/api/dealer/profile", "Dealer profile data"],
        ["PUT", "/api/dealer/profile", "Update dealer profile"],
        ["GET", "/api/stores", "Store list"],
    ]),
    ("Buyer", [
        ["GET", "/api/wishlist", "Buyer wishlist"],
        ["POST", "/api/wishlist", "Add to wishlist"],
        ["DELETE", "/api/wishlist/[id]", "Remove from wishlist"],
    ]),
    ("AI", [
        ["POST", "/api/ai/describe", "Generate vehicle description"],
        ["POST", "/api/ai/concierge", "AI concierge chat"],
        ["POST", "/api/ai/valuation", "Vehicle valuation"],
        ["POST", "/api/ai/sentiment", "Lead sentiment analysis"],
    ]),
    ("Payments", [
        ["POST", "/api/payments/create-order", "Create Razorpay order"],
        ["POST", "/api/payments/verify", "Verify payment signature"],
    ]),
]

for section_name, rows in api_sections:
    p = doc.add_paragraph()
    r = p.add_run(section_name)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLUE

    add_table(
        ["Method", "Endpoint", "Description"],
        rows,
        col_widths=[0.8, 2.5, 3.2],
    )
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════

add_divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CaroBest Platform Guide v2.0 | March 2026")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.italic = True

# Save
output_path = os.path.expanduser("~/Downloads/CaroBest_Platform_Guide.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
